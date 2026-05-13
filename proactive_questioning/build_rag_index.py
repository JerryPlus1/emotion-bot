#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
构建 RAG 知识库索引

从多个来源加载知识：
1. knowledge_example.json - JSON 格式知识库
2. external_knowledge_base/ - 外部文档目录（PDF/Word/TXT）

使用方法:
    python build_rag_index.py           # 构建索引
    python build_rag_index.py --watch  # 监控外部文件变化
"""

import os
import sys
import json
import pickle
from pathlib import Path
from typing import List, Dict, Tuple, Optional

import numpy as np
import torch

# 路径配置
SCRIPT_DIR = Path(__file__).resolve().parent

# 知识库目录（proactive_questioning 专用）
RAG_KB_DIR = SCRIPT_DIR / "rag_knowledge_base"
EXTERNAL_KB_DIR = SCRIPT_DIR / "external_knowledge_base"
BGE_MODEL_PATH = "/root/autodl-tmp/model/bge-m3"
INDEX_OUTPUT = RAG_KB_DIR / "knowledge_base.pkl"


def detect_json_text_field(json_path: Path) -> str:
    """自动检测 JSON 文件中的文本字段"""
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        if isinstance(data, list) and data:
            item = data[0]
            for field in ["content", "knowledge", "text", "description"]:
                if field in item and isinstance(item[field], str) and item[field]:
                    return field
        elif isinstance(data, dict):
            for field in ["content", "knowledge", "text", "description"]:
                if field in data and isinstance(data[field], str) and data[field]:
                    return field
    except Exception:
        pass
    return "content"


def load_json_documents(
    json_path: Path,
    text_field: Optional[str] = None,
) -> Tuple[List[str], List[Dict]]:
    """加载 JSON 格式文档"""
    if text_field is None:
        text_field = detect_json_text_field(json_path)
    
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    documents = []
    metadatas = []
    
    if isinstance(data, list):
        items = data
    else:
        items = data.get("documents", [data])
    
    for item in items:
        text = item.get(text_field, "")
        if not text:
            continue
        
        metadata = {
            "source": json_path.stem,
            "category": item.get("category", ""),
        }
        metadata.update({k: v for k, v in item.items() if k != text_field})
        
        documents.append(text)
        metadatas.append(metadata)
    
    return documents, metadatas


def load_text_document(
    file_path: Path,
) -> Tuple[List[str], List[Dict]]:
    """加载文本文件（按段落分割）"""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    
    # 按段落分割（双换行或单换行）
    paragraphs = []
    current = []
    
    for line in content.split("\n"):
        line = line.strip()
        if line:
            current.append(line)
        elif current:
            paragraph = " ".join(current)
            if len(paragraph) > 10:  # 过滤太短的段落
                paragraphs.append(paragraph)
            current = []
    
    if current:
        paragraph = " ".join(current)
        if len(paragraph) > 10:
            paragraphs.append(paragraph)
    
    metadatas = [
        {
            "source": file_path.stem,
            "source_type": "txt",
            "file": str(file_path.name),
        }
        for _ in paragraphs
    ]
    
    return paragraphs, metadatas


def load_pdf_document(
    file_path: Path,
) -> Tuple[List[str], List[Dict]]:
    """加载 PDF 文档"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print(f"  警告: PyMuPDF 未安装，跳过 PDF: {file_path.name}")
        return [], []
    
    documents = []
    metadatas = []
    
    try:
        with fitz.open(file_path) as pdf:
            for page_num, page in enumerate(pdf):
                text = page.get_text().strip()
                if text and len(text) > 50:
                    documents.append(text)
                    metadatas.append({
                        "source": file_path.stem,
                        "source_type": "pdf",
                        "file": str(file_path.name),
                        "page": page_num + 1,
                        "total_pages": len(pdf),
                    })
    except Exception as e:
        print(f"  警告: PDF 读取失败 {file_path.name}: {e}")
    
    return documents, metadatas


def load_docx_document(
    file_path: Path,
) -> Tuple[List[str], List[Dict]]:
    """加载 Word 文档"""
    try:
        from docx import Document
    except ImportError:
        print(f"  警告: python-docx 未安装，跳过 Word: {file_path.name}")
        return [], []
    
    documents = []
    metadatas = []
    
    try:
        doc = Document(file_path)
        
        # 收集所有段落
        current_para = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                current_para.append(text)
            elif current_para:
                full_text = " ".join(current_para)
                if len(full_text) > 50:
                    documents.append(full_text)
                    metadatas.append({
                        "source": file_path.stem,
                        "source_type": "docx",
                        "file": str(file_path.name),
                    })
                current_para = []
        
        # 处理最后一段
        if current_para:
            full_text = " ".join(current_para)
            if len(full_text) > 50:
                documents.append(full_text)
                metadatas.append({
                    "source": file_path.stem,
                    "source_type": "docx",
                    "file": str(file_path.name),
                })
        
        # 处理表格
        for t_idx, table in enumerate(doc.tables):
            table_texts = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)
            if table_texts:
                documents.append(" | ".join(table_texts))
                metadatas.append({
                    "source": file_path.stem,
                    "source_type": "docx_table",
                    "file": str(file_path.name),
                    "table_idx": t_idx,
                })
        
    except Exception as e:
        print(f"  警告: Word 读取失败 {file_path.name}: {e}")
    
    return documents, metadatas


def load_external_documents(
    directory: Path,
) -> Tuple[List[str], List[Dict]]:
    """从外部目录加载所有支持的文档"""
    if not directory.exists():
        print(f"  外部知识库目录不存在: {directory}")
        return [], []
    
    all_documents = []
    all_metadatas = []
    
    # 扫描所有文件
    for file_path in sorted(directory.rglob("*")):
        if not file_path.is_file():
            continue
        
        ext = file_path.suffix.lower()
        
        if ext == ".txt":
            print(f"  加载 TXT: {file_path.name}")
            docs, metas = load_text_document(file_path)
        elif ext == ".pdf":
            print(f"  加载 PDF: {file_path.name}")
            docs, metas = load_pdf_document(file_path)
        elif ext == ".docx":
            print(f"  加载 Word: {file_path.name}")
            docs, metas = load_docx_document(file_path)
        else:
            continue
        
        all_documents.extend(docs)
        all_metadatas.extend(metas)
        print(f"    -> {len(docs)} 个文本块")
    
    return all_documents, all_metadatas


def chunk_text(
    texts: List[str],
    metadatas: List[Dict],
    chunk_size: int = 512,
    overlap: int = 50,
) -> Tuple[List[str], List[Dict]]:
    """将长文本分割成小块"""
    chunks = []
    chunk_metas = []
    
    for text, meta in zip(texts, metadatas):
        if len(text) <= chunk_size:
            chunks.append(text)
            chunk_metas.append(meta)
        else:
            start = 0
            idx = 0
            while start < len(text):
                end = start + chunk_size
                chunk = text[start:end]
                chunks.append(chunk)
                chunk_metas.append({
                    **meta,
                    "chunk_idx": idx,
                    "chunk_start": start,
                    "chunk_end": min(end, len(text)),
                })
                start = end - overlap
                idx += 1
    
    return chunks, chunk_metas


def load_embedder():
    """加载嵌入模型（优先使用本地模型）"""
    from transformers import AutoTokenizer, AutoModel
    
    model_path = BGE_MODEL_PATH
    
    # 检查本地模型是否完整（需要 pytorch_model.bin 或 model.safetensors）
    has_model = os.path.exists(os.path.join(model_path, "pytorch_model.bin")) or \
                 os.path.exists(os.path.join(model_path, "model.safetensors"))
    
    if not has_model:
        # 使用更小的中文模型（自动下载，~500MB）
        print(f"本地模型不完整，使用轻量中文模型: paraphrase-multilingual-MiniLM-L12-v2")
        model_path = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
    else:
        print(f"使用本地模型: {model_path}")
    
    tokenizer = AutoTokenizer.from_pretrained(model_path, trust_remote_code=True)
    model = AutoModel.from_pretrained(
        model_path, trust_remote_code=True, torch_dtype=torch.float16
    ).cuda()
    model.eval()
    
    return tokenizer, model


def encode_texts(
    tokenizer, 
    model, 
    texts: List[str], 
    batch_size: int = 8,
    show_progress: bool = True,
):
    """批量编码文本"""
    all_embeddings = []
    total = len(texts)
    
    with torch.no_grad():
        for i in range(0, len(texts), batch_size):
            batch = texts[i : i + batch_size]
            inputs = tokenizer(
                batch,
                padding=True,
                truncation=True,
                max_length=512,
                return_tensors="pt",
            ).to(model.device)
            
            outputs = model(**inputs)
            hidden = outputs.last_hidden_state
            
            # Mean pooling
            mask = inputs["attention_mask"].unsqueeze(-1).expand(hidden.size()).float()
            sum_emb = torch.sum(hidden * mask, dim=1)
            sum_mask = torch.clamp(mask.sum(dim=1), min=1e-9)
            emb = sum_emb / sum_mask
            
            # L2 normalize
            emb = emb / emb.norm(dim=-1, keepdim=True)
            
            all_embeddings.append(emb.cpu().float().numpy())
            
            if show_progress and (i + batch_size) % 100 == 0:
                print(f"    编码进度: {i + batch_size}/{total} ({100*(i + batch_size)//total}%)")
    
    return np.concatenate(all_embeddings, axis=0)


def build_index():
    """构建 RAG 知识库索引"""
    print("=" * 60)
    print("RAG 知识库索引构建")
    print("=" * 60)
    
    all_documents = []
    all_metadatas = []
    
    # 1. 加载 JSON 知识库
    json_files = list(RAG_KB_DIR.glob("*.json")) + list(RAG_KB_DIR.glob("*.jsonl"))
    json_files = [f for f in json_files if f.name not in ["knowledge_base.pkl_chunks.json", "knowledge_chunks.json"]]
    
    if json_files:
        print("\n[1] 加载 JSON 知识库")
        for json_file in json_files:
            text_field = detect_json_text_field(json_file)
            print(f"  {json_file.name} (字段: {text_field})")
            docs, metas = load_json_documents(json_file, text_field)
            all_documents.extend(docs)
            all_metadatas.extend(metas)
            print(f"    -> {len(docs)} 条知识")
    
    # 2. 加载外部文档
    print(f"\n[2] 加载外部知识库: {EXTERNAL_KB_DIR}")
    ext_docs, ext_metas = load_external_documents(EXTERNAL_KB_DIR)
    
    if ext_docs:
        # 对长文本进行分块
        print(f"  分块处理 {len(ext_docs)} 个文档...")
        ext_docs, ext_metas = chunk_text(ext_docs, ext_metas, chunk_size=512, overlap=50)
        all_documents.extend(ext_docs)
        all_metadatas.extend(ext_metas)
        print(f"  外部知识库: {len(ext_docs)} 个文本块")
    else:
        # 创建目录并提示
        EXTERNAL_KB_DIR.mkdir(parents=True, exist_ok=True)
        print(f"  目录已创建: {EXTERNAL_KB_DIR}")
        print("  提示: 将 PDF/Word/TXT 文件放入该目录，重新运行即可加载")
    
    # 3. 检查是否有文档
    if not all_documents:
        print("\n错误: 没有找到任何知识库文档")
        print(f"请将文档放入以下目录:")
        print(f"  1. {RAG_KB_DIR} (JSON/JSONL 格式)")
        print(f"  2. {EXTERNAL_KB_DIR} (PDF/Word/TXT 格式)")
        return
    
    print(f"\n[3] 共 {len(all_documents)} 个文本块")
    
    # 4. 加载嵌入模型并编码
    print("\n[4] 加载嵌入模型...")
    tokenizer, model = load_embedder()
    embedding_dim = 1024  # BGE-M3 维度
    
    print(f"\n[5] 编码 {len(all_documents)} 个文本块...")
    vectors = encode_texts(tokenizer, model, all_documents, batch_size=16)
    print(f"    向量编码完成: {vectors.shape}")
    use_vector = True
    
    # 6. 保存索引
    print("\n[6] 保存索引...")
    RAG_KB_DIR.mkdir(parents=True, exist_ok=True)
    
    data = {
        "documents": all_documents,
        "metadatas": all_metadatas,
        "vectors": vectors,
        "embedding_dim": embedding_dim,
        "use_vector": use_vector,
    }
    
    with open(INDEX_OUTPUT, "wb") as f:
        pickle.dump(data, f)
    
    # 保存 chunks JSON
    chunks_path = RAG_KB_DIR / "knowledge_chunks.json"
    with open(chunks_path, "w", encoding="utf-8") as f:
        json.dump({"chunks": all_documents, "metadatas": all_metadatas}, f, ensure_ascii=False, indent=2)
    
    print(f"\n索引已保存: {INDEX_OUTPUT}")
    print("=" * 60)
    print("构建完成！")
    print("=" * 60)


def main():
    build_index()


if __name__ == "__main__":
    main()
